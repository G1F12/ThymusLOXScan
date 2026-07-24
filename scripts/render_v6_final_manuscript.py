"""Render the frozen v6 Markdown manuscript to HTML, PDF, and DOCX.

This is a presentation-only renderer. It does not execute any scientific code.
"""

from __future__ import annotations

import io
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "LOX_thymus_aging_public_preprint_v6_final.md"
CSS = ROOT / "manuscript" / "v6_final_preprint.css"
HTML = ROOT / "manuscript" / "pdf" / "LOX_thymus_aging_public_preprint_v6_final.html"
PDF = ROOT / "manuscript" / "pdf" / "LOX_thymus_aging_public_preprint_v6_final.pdf"
DOCX = ROOT / "manuscript" / "docx" / "LOX_thymus_aging_public_preprint_v6_final.docx"
PANDOC = Path(r"C:\Users\sasha\AppData\Local\Pandoc\pandoc.exe")
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")


def run(command: list[str]) -> None:
    subprocess.run(command, cwd=ROOT, check=True)


def add_page_field(paragraph) -> None:
    paragraph.alignment = 1
    run_obj = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run_obj._r.append(element)
    run_obj.font.name = "Arial"
    run_obj.font.size = Pt(9)
    run_obj.font.color.rgb = RGBColor(90, 90, 90)


def polish_docx() -> None:
    doc = Document(DOCX)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 22, RGBColor(23, 54, 93)),
        ("Heading 1", 15, RGBColor(23, 54, 93)),
        ("Heading 2", 12, RGBColor(36, 79, 120)),
    ):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
            style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.footer_distance = Mm(8)
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        add_page_field(paragraph)

    for table in doc.tables:
        table.autofit = True
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for row in table.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    props = doc.core_properties
    props.title = "Subtype-dependent LOX-family transcript changes in aging murine thymic stroma"
    props.author = "Aliaksandr Karatseyeu"
    props.subject = "ThymusLOXScan manuscript version 6.0"
    props.keywords = "thymus; aging; single-cell RNA-seq; lysyl oxidase; computational biology"
    doc.save(DOCX)


def stamp_pdf_page_numbers(raw_pdf: Path) -> None:
    reader = PdfReader(raw_pdf)
    writer = PdfWriter()
    for index, page in enumerate(reader.pages, start=1):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        overlay_bytes = io.BytesIO()
        c = canvas.Canvas(overlay_bytes, pagesize=(width, height))
        c.setFont("Helvetica", 9)
        c.setFillColorRGB(0.35, 0.35, 0.35)
        c.drawCentredString(width / 2, 15, str(index))
        c.save()
        overlay_bytes.seek(0)
        overlay = PdfReader(overlay_bytes).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": "Subtype-dependent LOX-family transcript changes in aging murine thymic stroma",
            "/Author": "Aliaksandr Karatseyeu",
            "/Subject": "ThymusLOXScan manuscript version 6.0",
        }
    )
    with PDF.open("wb") as handle:
        writer.write(handle)


def main() -> None:
    HTML.parent.mkdir(parents=True, exist_ok=True)
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    common = [
        str(PANDOC),
        str(SOURCE),
        "--from=gfm+fenced_divs",
        f"--resource-path={ROOT / 'manuscript'};{ROOT}",
        "--metadata",
        "lang=en",
    ]
    run(
        common
        + [
            "--standalone",
            "--embed-resources",
            f"--css={CSS}",
            "--output",
            str(HTML),
        ]
    )
    run(common + ["--output", str(DOCX)])
    polish_docx()

    with tempfile.TemporaryDirectory(prefix="thymusloxscan-v6-render-") as tmp:
        raw_pdf = Path(tmp) / "manuscript_raw.pdf"
        run(
            [
                str(EDGE),
                "--headless=new",
                "--disable-gpu",
                "--no-pdf-header-footer",
                f"--print-to-pdf={raw_pdf}",
                HTML.as_uri(),
            ]
        )
        if not raw_pdf.exists():
            raise RuntimeError("Edge did not create the expected PDF")
        stamp_pdf_page_numbers(raw_pdf)

    print(f"Rendered {HTML}")
    print(f"Rendered {PDF}")
    print(f"Rendered {DOCX}")


if __name__ == "__main__":
    main()
